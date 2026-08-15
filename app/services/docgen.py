"""docgen.py — rendering a structured contract into a real .docx file.

The only module that touches python-docx. It is a pure function of its input: the LLM
produces the *content* as JSON, this produces the *document*. Keeping generation and
rendering apart is what makes the output format deterministic — the model cannot
accidentally emit Markdown asterisks into a Word file, and this can be unit-tested
without any network access.
"""

import logging
import os
import re
import tempfile
from html import escape
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

logger = logging.getLogger(__name__)

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_MIME = "application/pdf"

# Cyrillic-capable and present on virtually every system that opens these files.
_BODY_FONT = "Times New Roman"
_BODY_SIZE = Pt(12)


def _safe_filename(title: str, fallback: str = "document") -> str:
    """A filesystem- and header-safe basename derived from the contract title."""
    name = re.sub(r"[^\w\s-]", "", title or "", flags=re.UNICODE).strip()
    name = re.sub(r"\s+", "_", name)
    return (name or fallback)[:60]


def _style_document(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = _BODY_FONT
    style.font.size = _BODY_SIZE
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)     # binding edge, per local convention
        section.right_margin = Cm(1.5)


def render_contract(contract: Dict[str, Any], output_dir: Optional[str] = None) -> str:
    """Renders a contract dict into a .docx and returns the local path.

    Expects the shape produced by `GeminiClient.draft_contract`:
        {title, intro?, sections: [{number, heading, body}], signature_blocks?: [...]}
    """
    doc = Document()
    _style_document(doc)

    title = (contract.get("title") or "Shartnoma").strip()
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()

    intro = (contract.get("intro") or "").strip()
    if intro:
        p = doc.add_paragraph(intro)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph()

    sections: List[Dict[str, Any]] = contract.get("sections") or []
    for idx, section in enumerate(sections, start=1):
        number = section.get("number") or idx
        sec_heading = (section.get("heading") or "").strip()
        if sec_heading:
            # Don't double-number when the model already prefixed the heading.
            label = sec_heading if re.match(r"^\s*\d+[.)]", sec_heading) else f"{number}. {sec_heading}"
            hp = doc.add_paragraph()
            hrun = hp.add_run(label)
            hrun.bold = True

        body = (section.get("body") or "").strip()
        for para in [b for b in body.split("\n") if b.strip()]:
            bp = doc.add_paragraph(para.strip())
            bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph()

    blocks = contract.get("signature_blocks") or []
    if blocks:
        doc.add_paragraph()
        table = doc.add_table(rows=1, cols=max(2, len(blocks)))
        cells = table.rows[0].cells
        for i, block in enumerate(blocks[: len(cells)]):
            cell = cells[i]
            cell.text = ""
            role = cell.paragraphs[0]
            rrun = role.add_run((block.get("party_role") or "").strip())
            rrun.bold = True
            cell.add_paragraph((block.get("party_name") or "").strip())
            cell.add_paragraph()
            cell.add_paragraph("_______________________")
            cell.add_paragraph("(imzo / подпись)")

    target_dir = output_dir or tempfile.mkdtemp(prefix="advoai_docgen_")
    os.makedirs(target_dir, exist_ok=True)
    path = os.path.join(target_dir, f"{_safe_filename(title)}.docx")
    doc.save(path)
    logger.info(f"Rendered contract '{title}' -> {path}")
    return path


_PDF_STYLE = """
    @page { size: A4; margin: 2cm 1.5cm 2cm 3cm; }
    body { font-family: "Noto Serif", "DejaVu Serif", Georgia, serif; font-size: 12pt;
           line-height: 1.5; color: #1a1a1a; }
    h1 { text-align: center; font-size: 15pt; letter-spacing: 0.02em; margin-bottom: 1.5em; }
    .intro { text-align: justify; margin-bottom: 1.5em; }
    .section-heading { font-weight: bold; margin-top: 1.2em; margin-bottom: 0.4em; }
    .section-body p { text-align: justify; margin: 0 0 0.6em 0; }
    table.signatures { width: 100%; border-collapse: collapse; margin-top: 2.5em; }
    table.signatures td { width: 50%; vertical-align: top; padding: 0 1em; }
    table.signatures .role { font-weight: bold; }
    table.signatures .line { margin-top: 2.5em; border-top: 1px solid #1a1a1a; width: 80%; }
    table.signatures .caption { font-size: 9pt; color: #555; margin-top: 0.2em; }
"""


def _contract_html(contract: Dict[str, Any]) -> str:
    """Renders the same section/clause model as `render_contract` to HTML for WeasyPrint.

    Kept a pure string-templating function (no Jinja) — the structure is simple enough
    that a template engine would add a dependency without buying anything.
    """
    title = (contract.get("title") or "Shartnoma").strip()
    parts = [f"<h1>{escape(title.upper())}</h1>"]

    intro = (contract.get("intro") or "").strip()
    if intro:
        parts.append(f'<p class="intro">{escape(intro)}</p>')

    sections: List[Dict[str, Any]] = contract.get("sections") or []
    for idx, section in enumerate(sections, start=1):
        number = section.get("number") or idx
        sec_heading = (section.get("heading") or "").strip()
        if sec_heading:
            label = sec_heading if re.match(r"^\s*\d+[.)]", sec_heading) else f"{number}. {sec_heading}"
            parts.append(f'<div class="section-heading">{escape(label)}</div>')

        body = (section.get("body") or "").strip()
        body_html = "".join(
            f"<p>{escape(p.strip())}</p>" for p in body.split("\n") if p.strip()
        )
        if body_html:
            parts.append(f'<div class="section-body">{body_html}</div>')

    blocks = contract.get("signature_blocks") or []
    if blocks:
        cells = "".join(
            f"""<td>
                <div class="role">{escape((b.get('party_role') or '').strip())}</div>
                <div>{escape((b.get('party_name') or '').strip())}</div>
                <div class="line"></div>
                <div class="caption">(imzo / подпись)</div>
            </td>"""
            for b in blocks
        )
        parts.append(f'<table class="signatures"><tr>{cells}</tr></table>')

    return f"""<!DOCTYPE html><html><head><meta charset="utf-8">
        <style>{_PDF_STYLE}</style></head><body>{"".join(parts)}</body></html>"""


def render_contract_pdf(contract: Dict[str, Any], output_dir: Optional[str] = None) -> str:
    """Renders a contract dict into a .pdf and returns the local path.

    Pure and network-free like `render_contract`, from the same structured contract
    JSON — no Chromium, no sidecar service, no network hop. Import is local to this
    function so a WeasyPrint import failure can't break DOCX generation, which has no
    dependency on it.
    """
    from weasyprint import HTML

    title = (contract.get("title") or "Shartnoma").strip()
    html = _contract_html(contract)

    target_dir = output_dir or tempfile.mkdtemp(prefix="advoai_docgen_")
    os.makedirs(target_dir, exist_ok=True)
    path = os.path.join(target_dir, f"{_safe_filename(title)}.pdf")
    HTML(string=html).write_pdf(path)
    logger.info(f"Rendered contract '{title}' -> {path}")
    return path


def contract_to_markdown(contract: Dict[str, Any]) -> str:
    """Flattens a contract to Markdown.

    Stored as the draft's `extracted_text` so the document goes into the same durable
    store as an upload — which is what lets the user say "change the salary" on the
    next turn and have the model actually see the draft it just produced.
    """
    lines = [f"# {contract.get('title', 'Shartnoma')}", ""]
    intro = (contract.get("intro") or "").strip()
    if intro:
        lines += [intro, ""]
    for idx, section in enumerate(contract.get("sections") or [], start=1):
        number = section.get("number") or idx
        lines.append(f"## {number}. {section.get('heading', '').strip()}")
        lines += [(section.get("body") or "").strip(), ""]
    for block in contract.get("signature_blocks") or []:
        lines.append(f"**{block.get('party_role', '')}**: {block.get('party_name', '')}")
    return "\n".join(lines)
